from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "manuals" / "MANUAL_BOOK_MAHASISWA_KP_FARMASI.docx"
ASSET_DIR = ROOT / "docs" / "manuals" / "assets" / "mahasiswa"


def set_run(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Manual Book Mahasiswa")
    set_run(run, 24, True, "0B6F7C")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SI-KP Farmasi UBP")
    set_run(run, 16, True, "0F172A")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Panduan penggunaan aplikasi Kerja Praktek untuk role Mahasiswa")
    set_run(run, 11, False, "475569")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Versi dokumentasi: 03 Juli 2026")
    set_run(run, 10, False, "64748B")
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    run.bold = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_steps(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_note(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    cell.paragraphs[0].add_run(title).bold = True
    cell.add_paragraph(body)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor.from_string("0F172A")
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, width: float = 6.2) -> None:
    if not path.exists():
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run(run, 9, False, "64748B")


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in [
        ("Heading 1", 16, "0B6F7C"),
        ("Heading 2", 13, "0B6F7C"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    add_title(doc)

    add_note(
        doc,
        "Tujuan panduan",
        "Dokumen ini membantu mahasiswa menjalankan alur Kerja Praktek di SI-KP Farmasi UBP mulai dari login sampai melihat nilai. "
        "Ikuti menu sesuai urutan proses KP dan perhatikan status yang tampil di setiap halaman.",
    )

    add_heading(doc, "1. Ringkasan Alur Mahasiswa")
    add_steps(
        doc,
        [
            "Masuk ke aplikasi melalui halaman login KP.",
            "Pastikan role aktif adalah Mahasiswa.",
            "Cek dan lengkapi Profil Saya bila masih diminta sistem.",
            "Buka Pendaftaran KP dan daftar pada periode yang sedang dibuka.",
            "Upload dokumen persyaratan di menu Berkas KP.",
            "Tunggu dokumen dan pendaftaran diverifikasi oleh Admin atau Koordinator KP.",
            "Saat jadwal pemilihan dibuka, pilih tempat KP di menu Pemilihan Tempat KP.",
            "Pantau Penempatan KP untuk melihat tempat, pembimbing dalam, dan preseptor.",
            "Isi Logbook KP selama pelaksanaan kerja praktek.",
            "Upload Laporan Akhir, ajukan Sidang bila sudah memenuhi syarat, lalu cek Nilai setelah dipublish.",
        ],
    )

    add_heading(doc, "2. Login")
    add_steps(
        doc,
        [
            "Buka alamat aplikasi KP.",
            "Masukkan email akun yang terdaftar di Core/KP.",
            "Masukkan kata sandi.",
            "Gunakan ikon mata untuk menampilkan atau menyembunyikan kata sandi bila perlu.",
            "Klik Buka Dashboard KP.",
        ],
    )
    add_note(
        doc,
        "Catatan login",
        "Jika akun punya lebih dari satu role, sistem akan menampilkan halaman pilih role. Pilih Mahasiswa untuk masuk ke dashboard mahasiswa.",
    )
    add_image(doc, ASSET_DIR / "01-login.png", "Screenshot 1. Halaman login SI-KP Farmasi UBP.")
    add_image(doc, ASSET_DIR / "03-login-password-toggle.png", "Screenshot 2. Form login dengan kontrol tampil/sembunyikan kata sandi.")

    add_heading(doc, "3. Dashboard Mahasiswa")
    add_bullets(
        doc,
        [
            "Dashboard menampilkan ringkasan status KP dan akses cepat ke modul utama.",
            "Gunakan tombol Profil untuk membuka Profil Saya.",
            "Gunakan menu di sidebar untuk pindah ke Pendaftaran KP, Berkas KP, Pemilihan Tempat KP, Penempatan KP, Logbook KP, Laporan Akhir, Sidang, dan Nilai.",
        ],
    )

    add_heading(doc, "4. Profil Saya")
    add_bullets(
        doc,
        [
            "Data identitas utama dibaca dari Core jika integrasi Core aktif.",
            "Data operasional KP yang memang khusus KP dapat dilengkapi di halaman profil.",
            "Jika nama, NIM, email, atau data resmi salah, perbaikan utama dilakukan di Core lalu refresh/sinkronkan data di KP.",
        ],
    )

    add_heading(doc, "5. Pendaftaran KP")
    add_steps(
        doc,
        [
            "Buka menu Pendaftaran KP.",
            "Pilih periode KP yang status pendaftarannya dibuka.",
            "Klik daftar atau lanjutkan pendaftaran.",
            "Setelah pendaftaran dibuat, lanjutkan ke Berkas KP untuk mengunggah dokumen.",
        ],
    )
    add_note(
        doc,
        "Status pendaftaran",
        "Draft berarti pendaftaran belum siap diverifikasi. Menunggu Verifikasi berarti dokumen wajib sudah lengkap dan sedang diperiksa. Terverifikasi berarti mahasiswa dapat lanjut ke tahap pemilihan tempat saat jadwal dibuka.",
    )

    add_heading(doc, "6. Berkas KP")
    add_steps(
        doc,
        [
            "Buka menu Berkas KP.",
            "Periksa daftar dokumen persyaratan pada periode aktif.",
            "Upload file sesuai format dan ukuran yang diminta sistem.",
            "Jika status dokumen Revisi, baca catatan revisi lalu upload ulang file yang benar.",
            "Jika semua dokumen wajib sudah disetujui, submit atau pantau status pendaftaran sampai menunggu verifikasi/terverifikasi.",
        ],
    )
    add_image(doc, ASSET_DIR / "02-berkas-kp-mobile.jpeg", "Screenshot 3. Halaman Berkas KP pada tampilan mobile.")

    add_heading(doc, "7. Pemilihan Tempat KP")
    add_bullets(
        doc,
        [
            "Menu ini hanya bisa digunakan jika pendaftaran sudah terverifikasi.",
            "Pemilihan tempat mengikuti jadwal yang ditetapkan pada periode KP.",
            "Waktu aplikasi mengikuti zona waktu Asia/Jakarta/WIB.",
            "Jika jadwal belum dibuka, mahasiswa menunggu sampai waktu pemilihan dimulai.",
            "Jika tempat sudah penuh, sistem dapat mengarahkan mahasiswa ke mekanisme daftar tunggu sesuai pengaturan.",
        ],
    )
    add_note(
        doc,
        "Penting",
        "Pilihan tempat KP bersifat terkunci untuk mahasiswa. Perubahan atau pembatalan hanya dapat dilakukan Admin atau Koordinator KP.",
    )

    add_heading(doc, "8. Penempatan KP")
    add_bullets(
        doc,
        [
            "Setelah pemilihan diproses, mahasiswa dapat melihat tempat KP resmi pada menu Penempatan KP.",
            "Halaman ini menampilkan tempat KP, pembimbing dalam, preseptor, dan status penempatan.",
            "Jika pembimbing belum tampil, tunggu Admin atau Koordinator KP menetapkan pembimbing.",
        ],
    )

    add_heading(doc, "9. Logbook KP")
    add_steps(
        doc,
        [
            "Buka menu Logbook KP.",
            "Tambahkan kegiatan sesuai tanggal pelaksanaan KP.",
            "Isi judul kegiatan, uraian kegiatan, dan bukti kegiatan bila diminta.",
            "Submit logbook untuk divalidasi pembimbing.",
            "Jika logbook direvisi, perbaiki sesuai catatan lalu submit ulang.",
        ],
    )

    add_heading(doc, "10. Laporan Akhir")
    add_bullets(
        doc,
        [
            "Upload laporan akhir setelah proses KP dan logbook siap.",
            "Jika laporan mendapat revisi, upload versi baru sesuai catatan pembimbing.",
            "Laporan akhir yang sudah disetujui menjadi syarat untuk pengajuan sidang sesuai aturan periode.",
        ],
    )

    add_heading(doc, "11. Sidang KP")
    add_steps(
        doc,
        [
            "Buka menu Sidang.",
            "Ajukan sidang jika laporan akhir sudah disetujui.",
            "Pantau jadwal sidang, ruang/link, pembimbing, dan penguji.",
            "Jika jadwal berubah, ikuti informasi terbaru pada halaman Sidang.",
        ],
    )

    add_heading(doc, "12. Nilai")
    add_bullets(
        doc,
        [
            "Nilai hanya tampil setelah dipublish oleh Admin atau Koordinator KP.",
            "Jika nilai belum tampil, artinya proses penilaian atau publikasi nilai belum selesai.",
            "Hubungi Admin Program jika nilai sudah diumumkan tetapi belum tampil di akun.",
        ],
    )

    add_heading(doc, "13. Kendala Umum dan Solusi")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Kendala"
    hdr[1].text = "Solusi"
    rows = [
        ("Tidak bisa login", "Pastikan email dan password benar. Jika akun baru dari Core, pastikan sudah punya akses aplikasi KP dan hubungi Admin jika masih gagal."),
        ("Role Mahasiswa tidak muncul", "Hubungi Admin Program untuk mengecek app access/role KP."),
        ("Tidak bisa daftar KP", "Cek apakah periode pendaftaran sudah dibuka dan profil sudah lengkap."),
        ("Tidak bisa memilih tempat", "Pastikan pendaftaran sudah terverifikasi dan jadwal pemilihan sedang dibuka."),
        ("Dokumen ditolak/revisi", "Baca catatan revisi pada Berkas KP, upload ulang file yang sesuai."),
        ("Nilai belum tampil", "Tunggu nilai dipublish oleh Admin/Koordinator KP."),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right

    add_heading(doc, "14. Kontak Bantuan")
    doc.add_paragraph(
        "Jika ada kendala data akun, role, pendaftaran, dokumen, pemilihan tempat, penempatan, logbook, sidang, atau nilai, mahasiswa menghubungi Admin Program/Koordinator KP Farmasi."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "SI-KP Farmasi UBP - Manual Book Mahasiswa"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

